from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.models import Document


def _safe_decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="ignore")


def _df_to_text(df: pd.DataFrame, max_rows: int = 250) -> str:
    if df.empty:
        return "Empty table."
    clipped = df.head(max_rows).copy()
    return clipped.to_markdown(index=False)


def load_uploaded_file(uploaded_file) -> Document:
    name = uploaded_file.name
    suffix = Path(name).suffix.lower()
    raw = uploaded_file.getvalue()
    doc_id = Path(name).stem.lower().replace(" ", "_").replace("-", "_")[:60]

    metadata = {"filename": name, "size_bytes": len(raw), "source": "upload"}

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages: list[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"\n\n[Page {idx}]\n{text}")
        text = "".join(pages).strip()
        metadata["pages"] = len(reader.pages)
    elif suffix == ".docx":
        doc = DocxDocument(io.BytesIO(raw))
        blocks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(blocks)
    elif suffix in {".csv"}:
        df = pd.read_csv(io.BytesIO(raw))
        text = _df_to_text(df)
        metadata["columns"] = list(df.columns)
        metadata["rows"] = int(len(df))
    elif suffix in {".xlsx", ".xls"}:
        sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
        parts: list[str] = []
        for sheet_name, df in sheets.items():
            parts.append(f"\n\n[Sheet: {sheet_name}]\n{_df_to_text(df)}")
        text = "".join(parts)
        metadata["sheets"] = list(sheets.keys())
    elif suffix == ".json":
        parsed = json.loads(_safe_decode(raw))
        text = json.dumps(parsed, indent=2, ensure_ascii=False)
    else:
        text = _safe_decode(raw)

    if not text.strip():
        text = f"No extractable text found in {name}."

    return Document(doc_id=doc_id, title=name, source_type=suffix.replace(".", "") or "text", text=text, metadata=metadata)


def load_local_path(path: Path) -> Document:
    suffix = path.suffix.lower()
    raw = path.read_bytes()

    class LocalUpload:
        def __init__(self, name: str, data: bytes):
            self.name = name
            self._data = data

        def getvalue(self) -> bytes:
            return self._data

    doc = load_uploaded_file(LocalUpload(path.name, raw))
    doc.metadata["source"] = "sample"
    doc.metadata["path"] = str(path)
    return doc


def load_local_folder(folder: Path) -> list[Document]:
    docs: list[Document] = []
    if not folder.exists():
        return docs
    allowed = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls", ".json"}
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in allowed:
            try:
                docs.append(load_local_path(path))
            except Exception as exc:
                docs.append(
                    Document(
                        doc_id=path.stem,
                        title=path.name,
                        source_type="error",
                        text=f"Could not load file: {exc}",
                        metadata={"source": "sample", "path": str(path), "error": str(exc)},
                    )
                )
    return docs


def parse_questionnaire_file(uploaded_file) -> list[str]:
    doc = load_uploaded_file(uploaded_file)
    suffix = Path(uploaded_file.name).suffix.lower()
    questions: list[str] = []
    if suffix in {".csv", ".xlsx", ".xls"}:
        raw = uploaded_file.getvalue()
        if suffix == ".csv":
            df = pd.read_csv(io.BytesIO(raw))
        else:
            first = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            df = next(iter(first.values()))
        lower_cols = {c.lower().strip(): c for c in df.columns}
        q_col = None
        for candidate in ["question", "questions", "security question", "requirement", "control question"]:
            if candidate in lower_cols:
                q_col = lower_cols[candidate]
                break
        if q_col is None and len(df.columns) > 0:
            q_col = df.columns[0]
        if q_col:
            questions = [str(x).strip() for x in df[q_col].dropna().tolist() if str(x).strip()]
    else:
        for line in doc.text.splitlines():
            clean = line.strip(" -\t0123456789.)")
            if clean.endswith("?") or len(clean.split()) > 5:
                questions.append(clean)
    return questions
